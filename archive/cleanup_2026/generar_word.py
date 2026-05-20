#!/usr/bin/env python3
"""
Generador de la Memoria TFG en formato Word (.docx)
Convierte memoria_tfg.md → Memoria_TFG_Joan_Roma.docx
con portada, índice, imágenes embebidas y formato profesional.

Uso:
    pip3 install python-docx Pillow
    python3 generar_word.py
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("❌ Falta python-docx. Instálalo con:")
    print("   pip3 install --break-system-packages python-docx Pillow")
    sys.exit(1)

# ── Configuración ──────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "memoria_tfg.md"
OUT_FILE = SCRIPT_DIR / "Memoria_TFG_Joan_Roma.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 12
LINE_SPACING = 1.15
UA_BLUE = RGBColor(0, 61, 121)  # Azul UA


def create_cover_page(doc):
    """Crea la portada del TFG."""
    # Espaciado superior
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Universidad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSITAT D'ALACANT")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.color.rgb = UA_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD DE ALICANTE")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = UA_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Escuela Politécnica Superior")
    run.font.name = FONT_NAME
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Grado en Ingeniería Informática")
    run.font.name = FONT_NAME
    run.font.size = Pt(13)

    # Espaciado
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRABAJO DE FIN DE GRADO")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = UA_BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Sistema de Trading Algorítmico Híbrido\n"
        "Basado en Inteligencia Artificial y\n"
        "Procesamiento de Lenguaje Natural"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.bold = True

    # Espaciado
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Datos
    data = [
        ("Autor:", "Joan Romà Llorca"),
        ("Tutor:", "José Ignacio Abreu Salas"),
        ("Curso Académico:", "2025–2026"),
        ("Fecha de entrega:", "Junio 2026"),
    ]
    for label, value in data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label} ")
        run_l.font.name = FONT_NAME
        run_l.font.size = Pt(12)
        run_l.bold = True
        run_v = p.add_run(value)
        run_v.font.name = FONT_NAME
        run_v.font.size = Pt(12)

    doc.add_page_break()


def set_style(doc):
    """Configura estilos base del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING

    # Configurar estilos de encabezado
    for level, size, color in [
        ('Heading 1', FONT_SIZE_H1, UA_BLUE),
        ('Heading 2', FONT_SIZE_H2, UA_BLUE),
        ('Heading 3', FONT_SIZE_H3, RGBColor(0, 0, 0)),
    ]:
        if level in doc.styles:
            s = doc.styles[level]
            s.font.name = FONT_NAME
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = color
            s.paragraph_format.space_before = Pt(18 if level == 'Heading 1' else 12)
            s.paragraph_format.space_after = Pt(6)
            if level == 'Heading 1':
                s.paragraph_format.page_break_before = True


def set_margins(doc):
    """Márgenes estándar para TFG."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def add_table(doc, headers, rows, caption=None):
    """Añade una tabla formateada."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.italic = True
        p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row - blue background
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003D79"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            # Zebra striping
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image(doc, filename, caption=None, width_inches=5.5):
    """Añade una imagen con pie de foto."""
    img_path = SCRIPT_DIR / filename
    if not img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Imagen no encontrada: {filename}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.italic = True
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.italic = True
        p.paragraph_format.space_after = Pt(12)


def add_code_block(doc, code, language="python"):
    """Añade un bloque de código con fondo gris."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Fondo gris via shading en el párrafo
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_and_build(doc):
    """Parse the markdown and build the Word document."""
    with open(MD_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    skip_metadata = True  # Skip the title/metadata block at top

    while i < len(lines):
        line = lines[i]

        # Skip the initial title and metadata (portada ya creada)
        if skip_metadata:
            if line.startswith("---") and i > 5:
                skip_metadata = False
                i += 1
                continue
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip Pandoc-style image width attributes and empty figure references
        if line.strip().startswith("!["):
            # Extract image info
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                caption = match.group(1)
                img_file = match.group(2).split("{")[0].strip()
                add_image(doc, img_file, caption)
            i += 1
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("##"):
            heading_text = line[2:].strip()
            # Skip duplicate metadata headings
            if heading_text in ["Resumen", "Abstract", "Índice de Figuras",
                                "Índice de Tablas", "Agradecimientos", "Bibliografía",
                                "Anexos"]:
                doc.add_heading(heading_text, level=1)
            elif heading_text.startswith("Capítulo"):
                doc.add_heading(heading_text, level=1)
            else:
                doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rules (section separators)
        if line.strip() == "---":
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            # Collect the table
            table_lines = []
            caption = None

            # Check if previous line was a caption (*Tabla X.X: ...*)
            if i > 0 and lines[i - 1].strip().startswith("*Tabla"):
                caption = lines[i - 1].strip().strip("*")

            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                # Parse headers
                headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
                # Skip separator line (line with ---)
                data_start = 1
                if re.match(r'^[\s|:-]+$', table_lines[1]):
                    data_start = 2

                rows = []
                for tl in table_lines[data_start:]:
                    cells = [c.strip() for c in tl.split("|") if c.strip()]
                    if cells:
                        rows.append(cells)

                if headers and rows:
                    add_table(doc, headers, rows, caption)
            continue

        # Table caption lines (already consumed above)
        if line.strip().startswith("*Tabla"):
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraphs
        # Handle bold/italic inline
        text = line.strip()

        # Numbered lists
        list_match = re.match(r'^(\d+)\.\s+(.+)', text)
        if list_match:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_run(p, list_match.group(2))
            i += 1
            continue

        # Bullet lists
        if text.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, text[2:])
            i += 1
            continue

        # LaTeX formulas → plain text
        if text.startswith("$$"):
            formula = text.strip("$").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_formatted_run(p, text)
        i += 1


def _add_formatted_run(paragraph, text):
    """Add text with basic markdown formatting (bold, italic, code)."""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_BODY)
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_BODY)
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part:
            run = paragraph.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_BODY)


def main():
    print("📝 Generando Memoria TFG en Word...")
    print(f"   Entrada: {MD_FILE}")
    print(f"   Salida:  {OUT_FILE}")

    doc = Document()

    # 1. Configurar estilos y márgenes
    set_style(doc)
    set_margins(doc)

    # 2. Portada
    create_cover_page(doc)

    # 3. Contenido desde Markdown
    parse_and_build(doc)

    # 4. Guardar
    doc.save(str(OUT_FILE))

    size_mb = OUT_FILE.stat().st_size / (1024 * 1024)
    print(f"\n✅ Memoria Word generada con éxito!")
    print(f"   Archivo: {OUT_FILE}")
    print(f"   Tamaño:  {size_mb:.1f} MB")


if __name__ == "__main__":
    main()
